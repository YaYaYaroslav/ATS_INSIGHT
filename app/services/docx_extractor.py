import docx


def extract_text_from_docx(path: str) -> str:
    """Витягує текст з .docx: параграфи + таблиці."""
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
